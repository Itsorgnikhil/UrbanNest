import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_hackathon_document():
    doc = Document()
    
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Helper to set shading color on table cells
    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper to set cell padding
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("UrbanNest Lifestyle Store")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(37, 99, 235) # Royal Blue

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("Mini Hackathon Project Submission — Technical Architecture & Team Contributions")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(92, 99, 94)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Team Contributions Section
    h1 = doc.add_heading("1. Team Members & Roles / Contributions", level=1)
    h1.runs[0].font.name = "Arial"
    h1.runs[0].font.color.rgb = RGBColor(31, 36, 33)

    p_intro = doc.add_paragraph("Each of our 3 team members contributed to specific areas of the UrbanNest Lifestyle Store platform as outlined below:")
    p_intro.runs[0].font.name = "Arial"
    p_intro.paragraph_format.space_after = Pt(12)

    # Table for Team Members
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Team Member", "Role & Module", "Enrollment No.", "Email Address"]
    col_widths = [Inches(1.8), Inches(1.8), Inches(1.5), Inches(1.9)]

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_shading(hdr_cells[idx], "2563EB") # Royal Blue Header
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[idx].paragraphs[0]
        p.runs[0].font.name = "Arial"
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    data = [
        ("Nakul Sharma", "N8N.io Integration & AI Chatbot Workflow", "12023002001149", "nakulsharma4767@gmail.com"),
        ("Nikhil Khumawat", "Deployment & GitHub Version Control", "12023002001138", "nikhilkmt123@gmail.com"),
        ("Devraj Singh Shekhawat", "UI/UX Design & Website Development", "12023002001148", "theraaj14@gmail.com")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F7F9F8" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = cell_value
            set_cell_shading(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.color.rgb = RGBColor(31, 36, 33)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 2. Technology Stack Section
    h2 = doc.add_heading("2. Technology Stack Overview", level=1)
    h2.runs[0].font.name = "Arial"
    h2.runs[0].font.color.rgb = RGBColor(31, 36, 33)

    tech_items = [
        ("Frontend Framework", "React 18 (Vite build engine) providing modular component architecture and fast HMR."),
        ("UI/UX Design System", "Vanilla CSS Custom Design System featuring Glassmorphic panels, Apple Liquid Glass sticky sidebars, Dribbble Bento Grid card layouts, and custom Google Fonts (Outfit & Plus Jakarta Sans)."),
        ("Icons & Media Assets", "Lucide React vector icon library and curated studio lifestyle photography."),
        ("Automation Engine", "N8N.io Cloud / Local Webhooks for handling real-time customer query form submissions and store notifications."),
        ("AI Agent & Chatbot", "N8N Chat Trigger + OpenAI LangChain Agent workflow supporting natural conversation, store knowledge lookup, and native iframe/API dual integration."),
        ("Version Control", "GitHub Repository: https://github.com/Itsorgnikhil/UrbanNest"),
        ("Hosting & Cloud CDN", "Render Static Site hosting with global CDN edge distribution, continuous auto-deployment from main branch, and managed SSL/TLS.")
    ]

    for category, desc in tech_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        r_bold = p.add_run(f"{category}: ")
        r_bold.font.name = "Arial"
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(37, 99, 235)
        
        r_desc = p.add_run(desc)
        r_desc.font.name = "Arial"
        r_desc.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 3. System Architecture & Workflow Flowchart Description
    h3 = doc.add_heading("3. System Architecture & N8N Workflow", level=1)
    h3.runs[0].font.name = "Arial"
    h3.runs[0].font.color.rgb = RGBColor(31, 36, 33)

    p_arch = doc.add_paragraph("The UrbanNest digital platform connects customer interactions seamlessly to automated workflows:")
    p_arch.runs[0].font.name = "Arial"

    workflow_steps = [
        "1. Web Front-End (Client Side): Customers explore product collections, search items, toggle wishlist, add items to cart, and fill the N8N Customer Query form or talk to NestBot.",
        "2. Query Form Workflow (N8N Webhook): Form submissions trigger a POST request to N8N webhook endpoint (/webhook/urbannest-query-form), formatted into structured JSON, logged to database/sheets, and confirmed with instant feedback.",
        "3. AI Chatbot Workflow (N8N Chat Trigger): User questions posted to N8N endpoint (/webhook/3ccde469-b8fd-4b0a-a429-2ee80a25218a/chat) process through NestBot AI Agent to answer queries about products, timings (10 AM - 9 PM), location (Indiranagar), and delivery.",
        "4. Deployment Pipeline: Git commits pushed to GitHub repository automatically trigger Render CDN static build, updating live production seamlessly."
    ]

    for step in workflow_steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(step)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 4. Links & Deliverables
    h4 = doc.add_heading("4. Key Links & Deliverables", level=1)
    h4.runs[0].font.name = "Arial"
    h4.runs[0].font.color.rgb = RGBColor(31, 36, 33)

    links = [
        ("GitHub Repository", "https://github.com/Itsorgnikhil/UrbanNest"),
        ("Live Web App (Render)", "https://urbannest-lifestyle-store.onrender.com"),
        ("N8N Chatbot Endpoint", "https://nakul1122.app.n8n.cloud/webhook/3ccde469-b8fd-4b0a-a429-2ee80a25218a/chat"),
        ("N8N Query Form Workflow", "n8n_query_form_workflow.json (Included in project root)"),
        ("N8N Chatbot Workflow", "n8n_chatbot_workflow.json (Included in project root)")
    ]

    for label, url in links:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.font.name = "Arial"
        r_lbl.font.bold = True
        r_val = p.add_run(url)
        r_val.font.name = "Arial"
        r_val.font.color.rgb = RGBColor(37, 99, 235)

    doc.save("UrbanNest_Project_Submission.docx")
    print("Successfully created UrbanNest_Project_Submission.docx")

if __name__ == "__main__":
    create_hackathon_document()
